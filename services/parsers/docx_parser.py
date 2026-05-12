import io
from docx import Document
from services.parsers.txt_parser import parse_txt_content

def parse_docx_content(file_bytes: bytes, filename: str = "", options_count: int = None) -> list:
    """
    Видобування тексту з DOCX документів.
    """
    if filename.lower().endswith('.doc'):
        raise ValueError(
            "Старий формат '.doc' не підтримується системою. "
            "Будь ласка, відкрийте цей файл у Word, натисніть 'Зберегти як' "
            "та оберіть сучасний формат '.docx', після чого завантажте його знову."
        )

    docx_file = io.BytesIO(file_bytes)
    
    try:
        doc = Document(docx_file)
        text_blocks = []

        # 1. Збираємо текст з абзаців
        for paragraph in doc.paragraphs:
            # ВАЖЛИВО: Ми більше НЕ ігноруємо пусті рядки!
            # Ми просто очищаємо їх від пробілів. Якщо рядок був пустим, він стане ""
            # Це дозволить нам зберегти візуальні розриви (\n\n) між тестами.
            text_blocks.append(paragraph.text.strip())

        # 2. Збираємо текст з таблиць
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_blocks.append(" ".join(row_text))

        full_text = "\n".join(text_blocks)
        
        return parse_txt_content(full_text, options_count)
        
    except Exception as e:
        raise Exception(f"Помилка структури DOCX документа: {str(e)}")